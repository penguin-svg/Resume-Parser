import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client=Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"


job_description="""
Skill required: Tech for Operations - Artificial Intelligence (AI)
Designation: AI/ML Computational Science Associate
Qualifications:BE/BTech/MCA
Years of Experience:0 to 1 year
Language - Ability:English(International) - Intermediate
About Accenture
Accenture is a global professional services company with leading capabilities in digital, cloud and security.Combining unmatched experience and specialized skills across more than 40 industries, we offer Strategy and Consulting, Technology and Operations services, and Accenture Song— all powered by the world’s largest network of Advanced Technology and Intelligent Operations centers. Our 784,000 people deliver on the promise of technology and human ingenuity every day, serving clients in more than 120 countries. We embrace the power of change to create value and shared success for our clients, people, shareholders, partners and communities.Visit us at www.accenture.com
What would you do? About Accenture Operations Accenture Operations is a global leader in transforming business processes through technology, data, and human ingenuity. We help organizations move from transactional to transformational operations by leveraging advanced analytics, automation, and AI-powered platforms like SynOps. With over 220,000 professionals across 50+ delivery locations, Accenture Operations enables clients to access new performance frontiers, improve agility, and deliver exceptional experiences for customers and employees. Our approach creates 360° value—driving efficiency, innovation, and sustainable growth across industries. Job Summary We are looking for enthusiastic and highly motivated fresh graduates with a strong foundation in Python programming and software development concepts. The candidate will work with experienced development teams to design, develop, test, and maintain applications while gaining exposure to cloud, data engineering, AI, and modern software development practices. This role provides an excellent opportunity to build a career in application development, automation, data processing, and emerging technologies using Python. Experience 0–1 Years of Experience Freshers and recent graduates are encouraged to apply. Internship or academic project experience in Python is preferred. Qualifications B.E./B.Tech, MCA, M.Sc. (Computer Science) Strong academic background. Relevant certifications in Python, Cloud, Data Analytics, or AI are an added advantage.
What are we looking for? Good knowledge of Python programming. Understanding of Object-Oriented Programming (OOP) concepts. Basic knowledge of: Data Structures & Algorithms SQL and Database Concepts REST APIs Git/Version Control Software Development Life Cycle (SDLC) •Sample Projects (Good to Have) Python Automation Tools Web Applications using Django/Flask Data Analysis and Visualization Projects Chatbot or AI-based Applications Database-Driven Applications API Development Projects Exposure to: Django / Flask / FastAPI Pandas, NumPy HTML, CSS, JavaScript Cloud Platforms (Azure/AWS/GCP) AI/ML concepts Automation scripting
Roles and Responsibilities: Develop, test, and maintain Python-based applications and scripts. Participate in coding, debugging, and troubleshooting activities. Assist in developing REST APIs and backend services. Support automation and process improvement initiatives. Work on data processing, reporting, and analytics solutions. Participate in code reviews and follow coding standards. Collaborate with cross-functional teams to understand requirements and deliver solutions. Prepare technical documentation and project updates. Learn and adopt emerging technologies, tools, and frameworks. •Preferred Competencies Strong analytical and problem-solving skills. Excellent communication and teamwork skills. Eagerness to learn new technologies. Ability to work in a collaborative and fast-paced environment. Positive attitude and customer-centric mindset.
"""
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system={
    "role" : "system",
    "content" : system_prompt
}
message_user={
    "role" : "user",
    "content" : user_prompt
}
response_format={
    "type" : "json_object"
}


messages=[message_system, message_user]

response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)


answer=response.choices[0].message.content

raw_json=answer
# print(raw_json)



import json
job_data=json.loads(raw_json)

job = JobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)



#parse real
class MatchResult(BaseModel):
    score: float
    details: dict
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None



# lets do it now
resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    #C:\Users\Pratyush\padho_with_pratyush\week1\day5\resumes\abhay resume new - Abhay Singh.pdf
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    time.sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    #score and details
    #acount chtgpt
    # request bhejna shhur krega millions
    #chattgot server jam ho jayega
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2 = all_results[-2:]


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])